"""Pipeline ML del artículo: EDA, train 5 modelos, CV, tuning, McNemar, reportes."""

from __future__ import annotations

import json
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import joblib
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from catboost import CatBoostClassifier
from imblearn.over_sampling import SMOTE
from lightgbm import LGBMClassifier
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sklearn.ensemble import RandomForestClassifier, StackingClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    average_precision_score,
    confusion_matrix,
    f1_score,
    matthews_corrcoef,
    precision_score,
    recall_score,
    roc_auc_score,
    roc_curve,
    precision_recall_curve,
)
from sklearn.model_selection import RandomizedSearchCV, StratifiedKFold, cross_validate
from sklearn.neural_network import MLPClassifier
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.svm import SVC
from xgboost import XGBClassifier
from docx import Document

from app.config import ARTIFACTS_DIR, FIGURES_DIR, REPORTS_DIR
from app.ml.dataset import FEATURE_COLS, ensure_dataset, load_dataset, save_mock_dataset

RANDOM_STATE = 42
METRICS_PATH = ARTIFACTS_DIR / "metrics.json"
BEST_MODEL_PATH = ARTIFACTS_DIR / "best_model.joblib"
SPLIT_PATH = ARTIFACTS_DIR / "split_info.json"


def _ensure_data() -> pd.DataFrame:
    return load_dataset()


def temporal_split(df: pd.DataFrame, test_months: int = 1):
    """
    Split temporal usando TODOS los meses de la cohorte.
    - No elimina meses del dataset (EDA/ranking ven el periodo completo).
    - El test usa el/los últimos meses que tengan al menos un positivo (horizonte observable).
    - Meses finales sin y=1 (p.ej. julio incompleto) quedan fuera del par train/test,
      pero permanecen en la cohorte para EDA.
    """
    work = df.copy()
    work["cutoff_date"] = pd.to_datetime(work["cutoff_date"])
    dates = sorted(work["cutoff_date"].unique())

    if len(dates) <= 1:
        from sklearn.model_selection import train_test_split

        train, test = train_test_split(
            work, test_size=0.25, stratify=work["y"], random_state=RANDOM_STATE
        )
        return train, test, {
            "cut": "random_stratified_25pct",
            "months_in_cohort": [str(d) for d in dates],
            "mode": "fallback_stratified",
        }

    months_with_pos = [
        d for d in dates if int(work.loc[work["cutoff_date"] == d, "y"].sum()) > 0
    ]
    months_without_pos = [d for d in dates if d not in months_with_pos]

    if not months_with_pos:
        from sklearn.model_selection import train_test_split

        train, test = train_test_split(
            work, test_size=0.25, stratify=work["y"], random_state=RANDOM_STATE
        )
        return train, test, {
            "cut": "random_stratified_25pct",
            "months_in_cohort": [str(d) for d in dates],
            "mode": "fallback_stratified",
        }

    k = max(1, min(test_months, len(months_with_pos) - 1)) if len(months_with_pos) > 1 else 1
    # Test = últimos k meses CON positivos; train = meses anteriores con/sin positivo
    test_month_set = set(months_with_pos[-k:])
    cut = min(test_month_set)
    # Excluir del train/test solo meses posteriores al último mes con positivo (censurados)
    last_labeled = months_with_pos[-1]
    labeled = work[work["cutoff_date"] <= last_labeled].copy()
    train = labeled[labeled["cutoff_date"] < cut].copy()
    test = labeled[labeled["cutoff_date"] >= cut].copy()

    if train.empty or train["y"].nunique() < 2 or test.empty or int(test["y"].sum()) == 0:
        from sklearn.model_selection import train_test_split

        train, test = train_test_split(
            labeled, test_size=0.25, stratify=labeled["y"], random_state=RANDOM_STATE
        )
        return train, test, {
            "cut": "random_stratified_25pct",
            "months_in_cohort": [str(d) for d in dates],
            "months_with_positives": [str(d) for d in months_with_pos],
            "months_sin_positivo_horizonte": [str(d) for d in months_without_pos],
            "mode": "fallback_stratified",
        }

    return train, test, {
        "cut": str(cut),
        "months_in_cohort": [str(d) for d in dates],
        "months_with_positives": [str(d) for d in months_with_pos],
        "months_sin_positivo_horizonte": [str(d) for d in months_without_pos],
        "mode": "temporal",
        "test_months": k,
        "n_months_total": len(dates),
    }


def find_best_threshold(
    y_true,
    y_prob,
    metric: str = "mcc",
) -> dict[str, float]:
    """Elige umbral en [0.05, 0.95] que maximiza MCC (o F1). Se calibra en train."""
    y_true = np.asarray(y_true).astype(int)
    y_prob = np.asarray(y_prob, dtype=float)
    if len(np.unique(y_true)) < 2:
        return {"threshold": 0.5, "mcc": 0.0, "f1": 0.0, "metric": metric}

    grid = np.unique(
        np.concatenate(
            [
                np.linspace(0.05, 0.95, 37),
                np.quantile(y_prob, np.linspace(0.05, 0.95, 19)),
            ]
        )
    )
    best_t, best_score, best_mcc, best_f1 = 0.5, -1.0, -1.0, -1.0
    for t in grid:
        pred = (y_prob >= t).astype(int)
        if pred.min() == pred.max():
            continue
        mcc = float(matthews_corrcoef(y_true, pred))
        f1 = float(f1_score(y_true, pred, zero_division=0))
        score = mcc if metric == "mcc" else f1
        if score > best_score:
            best_t, best_score, best_mcc, best_f1 = float(t), float(score), mcc, f1

    return {
        "threshold": round(best_t, 4),
        "mcc": round(best_mcc, 4),
        "f1": round(best_f1, 4),
        "metric": metric,
    }


def build_models(scale_pos_weight: float) -> dict[str, Any]:
    m1 = LogisticRegression(max_iter=1000, class_weight="balanced", random_state=RANDOM_STATE)
    m2 = RandomForestClassifier(
        n_estimators=200, class_weight="balanced", random_state=RANDOM_STATE, n_jobs=-1
    )
    m3 = XGBClassifier(
        n_estimators=200,
        max_depth=4,
        learning_rate=0.08,
        subsample=0.9,
        colsample_bytree=0.9,
        scale_pos_weight=scale_pos_weight,
        eval_metric="logloss",
        random_state=RANDOM_STATE,
        n_jobs=-1,
    )
    h1 = StackingClassifier(
        estimators=[
            (
                "lgbm",
                LGBMClassifier(
                    n_estimators=150,
                    learning_rate=0.08,
                    class_weight="balanced",
                    random_state=RANDOM_STATE,
                    verbose=-1,
                ),
            ),
            (
                "cat",
                CatBoostClassifier(
                    iterations=150,
                    depth=4,
                    learning_rate=0.08,
                    auto_class_weights="Balanced",
                    verbose=0,
                    random_seed=RANDOM_STATE,
                ),
            ),
        ],
        final_estimator=LogisticRegression(max_iter=500, random_state=RANDOM_STATE),
        cv=3,
        n_jobs=1,
    )
    h2 = StackingClassifier(
        estimators=[
            (
                "mlp",
                Pipeline(
                    [
                        ("scaler", StandardScaler()),
                        (
                            "clf",
                            MLPClassifier(
                                hidden_layer_sizes=(64, 32),
                                max_iter=200,
                                random_state=RANDOM_STATE,
                            ),
                        ),
                    ]
                ),
            ),
            (
                "svm",
                Pipeline(
                    [
                        ("scaler", StandardScaler()),
                        (
                            "clf",
                            SVC(
                                kernel="rbf",
                                probability=True,
                                class_weight="balanced",
                                random_state=RANDOM_STATE,
                            ),
                        ),
                    ]
                ),
            ),
        ],
        final_estimator=LogisticRegression(max_iter=500, random_state=RANDOM_STATE),
        cv=3,
        n_jobs=1,
    )
    return {
        "M1_LogisticRegression": m1,
        "M2_RandomForest": m2,
        "M3_XGBoost": m3,
        "H1_LightGBM_CatBoost": h1,
        "H2_MLP_SVM": h2,
    }


def _scores(y_true, y_pred, y_prob) -> dict[str, float]:
    return {
        "precision": float(precision_score(y_true, y_pred, zero_division=0)),
        "recall": float(recall_score(y_true, y_pred, zero_division=0)),
        "f1": float(f1_score(y_true, y_pred, zero_division=0)),
        "accuracy": float(accuracy_score(y_true, y_pred)),
        "roc_auc": float(roc_auc_score(y_true, y_prob)) if len(np.unique(y_true)) > 1 else 0.5,
        "pr_auc": float(average_precision_score(y_true, y_prob))
        if len(np.unique(y_true)) > 1
        else float(np.mean(y_true)),
        "mcc": float(matthews_corrcoef(y_true, y_pred)),
    }


def _save_confusion(y_true, y_pred, name: str) -> str:
    cm = confusion_matrix(y_true, y_pred, labels=[0, 1])
    fig, ax = plt.subplots(figsize=(4.5, 3.8))
    sns.heatmap(
        cm,
        annot=True,
        fmt="d",
        cmap="Blues",
        ax=ax,
        xticklabels=["Sin falla", "Falla"],
        yticklabels=["Sin falla", "Falla"],
    )
    ax.set_xlabel("Predicho")
    ax.set_ylabel("Real")
    short = name.replace("_", " ")
    ax.set_title(f"Matriz de confusión — {short}")
    path = FIGURES_DIR / f"cm_{name}.png"
    fig.tight_layout()
    fig.savefig(path, dpi=140)
    plt.close(fig)
    return path.name


def _save_accuracy_comparison(results: dict[str, Any]) -> str:
    """Bar chart tipo Figura 3: accuracy de todos los modelos."""
    names = list(results.keys())
    accs = [float(results[n]["accuracy"]) for n in names]
    short = [n.split("_", 1)[0] if n.startswith(("M", "H")) else n for n in names]
    # etiquetas legibles
    labels = []
    for n in names:
        if n.startswith("M1"):
            labels.append("M1 LR")
        elif n.startswith("M2"):
            labels.append("M2 RF")
        elif n.startswith("M3"):
            labels.append("M3 XGB")
        elif n.startswith("H1"):
            labels.append("H1 LGBM+CB")
        elif n.startswith("H2"):
            labels.append("H2 MLP+SVM")
        else:
            labels.append(n)

    colors = ["#22c55e", "#ef4444", "#eab308", "#3b82f6", "#a855f7"]
    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    bars = ax.bar(labels, accs, color=colors[: len(labels)], edgecolor="white", width=0.65)
    ax.set_ylim(0, 1.05)
    ax.set_ylabel("Accuracy general")
    ax.set_title("Comparación de Accuracy entre Modelos")
    ax.axhline(0.5, color="#94a3b8", linestyle="--", linewidth=0.8, alpha=0.7)
    for bar, acc in zip(bars, accs):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            acc + 0.02,
            f"{acc:.2f}",
            ha="center",
            va="bottom",
            fontsize=10,
            fontweight="bold",
        )
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    path = FIGURES_DIR / "accuracy_comparison.png"
    fig.tight_layout()
    fig.savefig(path, dpi=140)
    plt.close(fig)
    return path.name


def regenerate_analysis_figures_from_metrics() -> dict[str, str] | None:
    """Regenera gráfico de accuracy desde metrics.json (sin reentrenar)."""
    if not METRICS_PATH.exists():
        return None
    payload = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    models = payload.get("models") or {}
    if not models:
        return None
    acc_file = _save_accuracy_comparison(models)
    payload["figures"] = {**(payload.get("figures") or {}), "accuracy_comparison": acc_file}
    METRICS_PATH.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return payload["figures"]


def _save_roc(y_true, y_prob, name: str) -> str:
    fpr, tpr, _ = roc_curve(y_true, y_prob)
    fig, ax = plt.subplots(figsize=(4.5, 3.8))
    ax.plot(fpr, tpr, label=name)
    ax.plot([0, 1], [0, 1], "--", color="gray")
    ax.set_xlabel("FPR")
    ax.set_ylabel("TPR")
    ax.set_title(f"Curva ROC — {name}")
    ax.legend(loc="lower right")
    path = FIGURES_DIR / f"roc_{name}.png"
    fig.tight_layout()
    fig.savefig(path, dpi=120)
    plt.close(fig)
    return path.name


def _save_pr(y_true, y_prob, name: str) -> str:
    p, r, _ = precision_recall_curve(y_true, y_prob)
    fig, ax = plt.subplots(figsize=(4.5, 3.8))
    ax.plot(r, p, label=name)
    ax.set_xlabel("Recall")
    ax.set_ylabel("Precision")
    ax.set_title(f"Curva PR — {name}")
    ax.legend(loc="lower left")
    path = FIGURES_DIR / f"pr_{name}.png"
    fig.tight_layout()
    fig.savefig(path, dpi=120)
    plt.close(fig)
    return path.name


def _save_corr_heatmap(df: pd.DataFrame) -> str:
    cols = [c for c in FEATURE_COLS if c in df.columns] + ["y"]
    corr = df[cols].corr(numeric_only=True)
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.heatmap(corr, cmap="coolwarm", center=0, ax=ax)
    ax.set_title("Mapa de calor — correlaciones")
    path = FIGURES_DIR / "heatmap_corr.png"
    fig.tight_layout()
    fig.savefig(path, dpi=120)
    plt.close(fig)
    return path.name


def mcnemar_test(y_true, pred_a, pred_b) -> dict[str, float]:
    # tabla de discordancia
    b = int(np.sum((pred_a == y_true) & (pred_b != y_true)))
    c = int(np.sum((pred_a != y_true) & (pred_b == y_true)))
    if b + c == 0:
        return {"chi2": 0.0, "p_value": 1.0, "b": b, "c": c}
    chi2 = (abs(b - c) - 1) ** 2 / (b + c)
    # aproximación p con chi2 df=1
    from scipy.stats import chi2 as chi2_dist

    p = float(1 - chi2_dist.cdf(chi2, df=1))
    return {"chi2": float(chi2), "p_value": p, "b": b, "c": c}


def eda_summary() -> dict[str, Any]:
    df = _ensure_data()
    heatmap = _save_corr_heatmap(df)
    by_month = (
        df.groupby(df["cutoff_date"].dt.strftime("%Y-%m"))
        .agg(n=("y", "size"), positivos=("y", "sum"))
        .reset_index()
        .rename(columns={"cutoff_date": "mes"})
    )
    by_month["prevalencia_pct"] = (100 * by_month["positivos"] / by_month["n"]).round(2)
    return {
        "n_observaciones": int(len(df)),
        "n_vehiculos": int(df["vehicle_id"].nunique()),
        "n_features": len(FEATURE_COLS),
        "periodo": {
            "inicio": str(df["cutoff_date"].min().date()),
            "fin": str(df["cutoff_date"].max().date()),
        },
        "prevalencia_global_pct": round(100 * float(df["y"].mean()), 2),
        "tabla_clases": [
            {"clase": "Sin falla correctiva (y=0)", "n": int((df["y"] == 0).sum())},
            {"clase": "Falla correctiva 30d (y=1)", "n": int((df["y"] == 1).sum())},
        ],
        "por_mes": by_month.to_dict(orient="records"),
        "descriptivos": df[FEATURE_COLS].describe().round(3).fillna(0).to_dict(),
        "figuras": {"heatmap": heatmap},
        "features": FEATURE_COLS,
    }


def run_full_pipeline(
    n_folds: int = 5,
    do_tuning: bool = True,
) -> dict[str, Any]:
    ensure_dataset()
    df = _ensure_data()
    train_df, test_df, split_meta = temporal_split(df, test_months=1)
    cut = split_meta["cut"] if isinstance(split_meta, dict) else str(split_meta)
    X_train = train_df[FEATURE_COLS]
    y_train = train_df["y"].astype(int)
    X_test = test_df[FEATURE_COLS]
    y_test = test_df["y"].astype(int)

    smote = SMOTE(random_state=RANDOM_STATE)
    X_res, y_res = smote.fit_resample(X_train, y_train)
    pos = max(1, int(y_res.sum()))
    neg = max(1, int(len(y_res) - pos))
    spw = neg / pos

    models = build_models(spw)
    results: dict[str, Any] = {}
    preds: dict[str, np.ndarray] = {}
    probs: dict[str, np.ndarray] = {}
    fitted: dict[str, Any] = {}
    thresholds: dict[str, float] = {}

    for name, model in models.items():
        t0 = time.perf_counter()
        model.fit(X_res, y_res)
        elapsed = time.perf_counter() - t0

        # Umbral óptimo calibrado en train (sin SMOTE) → se aplica en test
        y_prob_train = model.predict_proba(X_train)[:, 1]
        thr_info = find_best_threshold(y_train, y_prob_train, metric="mcc")
        thr = float(thr_info["threshold"])
        thresholds[name] = thr

        y_prob = model.predict_proba(X_test)[:, 1]
        y_pred = (y_prob >= thr).astype(int)
        metrics = _scores(y_test, y_pred, y_prob)
        metrics["train_seconds"] = round(elapsed, 3)
        metrics["threshold"] = thr
        metrics["threshold_train_mcc"] = thr_info["mcc"]
        metrics["threshold_train_f1"] = thr_info["f1"]
        cm_file = _save_confusion(y_test, y_pred, name)
        roc_file = _save_roc(y_test, y_prob, name)
        pr_file = _save_pr(y_test, y_prob, name)
        results[name] = {
            **metrics,
            "figures": {"confusion": cm_file, "roc": roc_file, "pr": pr_file},
        }
        preds[name] = y_pred
        probs[name] = y_prob
        fitted[name] = model

    # mejor por PR-AUC luego MCC (métricas ya con umbral óptimo)
    best_name = max(results.keys(), key=lambda k: (results[k]["pr_auc"], results[k]["mcc"]))
    best_model = fitted[best_name]
    best_threshold = thresholds[best_name]

    # CV sobre resampled train
    scoring = {
        "f1": "f1",
        "roc_auc": "roc_auc",
        "accuracy": "accuracy",
    }
    y_res_arr = np.asarray(y_res)
    min_class = int(np.bincount(y_res_arr).min()) if len(y_res_arr) else 2
    n_folds_eff = min(n_folds, max(2, min_class))
    cv = StratifiedKFold(n_splits=n_folds_eff, shuffle=True, random_state=RANDOM_STATE)
    cv_table = {}
    for name, model in build_models(spw).items():
        # re-instanciar para CV limpio
        scores = cross_validate(model, X_res, y_res, cv=cv, scoring=scoring, n_jobs=-1)
        cv_table[name] = {
            "f1_mean": float(scores["test_f1"].mean()),
            "f1_std": float(scores["test_f1"].std()),
            "roc_auc_mean": float(scores["test_roc_auc"].mean()),
            "roc_auc_std": float(scores["test_roc_auc"].std()),
            "accuracy_mean": float(scores["test_accuracy"].mean()),
            "accuracy_std": float(scores["test_accuracy"].std()),
        }

    tuning = None
    if do_tuning and best_name.startswith("M"):
        # tuning ligero solo en modelos base
        search_space = {
            "M1_LogisticRegression": (
                LogisticRegression(max_iter=1000, class_weight="balanced", random_state=RANDOM_STATE),
                {"C": [0.1, 0.5, 1.0, 2.0]},
            ),
            "M2_RandomForest": (
                RandomForestClassifier(class_weight="balanced", random_state=RANDOM_STATE, n_jobs=-1),
                {"n_estimators": [100, 200, 300], "max_depth": [None, 6, 10]},
            ),
            "M3_XGBoost": (
                XGBClassifier(
                    eval_metric="logloss",
                    scale_pos_weight=spw,
                    random_state=RANDOM_STATE,
                    n_jobs=-1,
                ),
                {
                    "n_estimators": [100, 200],
                    "max_depth": [3, 4, 6],
                    "learning_rate": [0.05, 0.08, 0.1],
                },
            ),
        }
        if best_name in search_space:
            base, grid = search_space[best_name]
            search = RandomizedSearchCV(
                base,
                grid,
                n_iter=6,
                cv=3,
                scoring="average_precision",
                random_state=RANDOM_STATE,
                n_jobs=-1,
            )
            search.fit(X_res, y_res)
            tuned = search.best_estimator_
            y_prob_train = tuned.predict_proba(X_train)[:, 1]
            thr_info = find_best_threshold(y_train, y_prob_train, metric="mcc")
            thr = float(thr_info["threshold"])
            y_prob = tuned.predict_proba(X_test)[:, 1]
            y_pred = (y_prob >= thr).astype(int)
            tuned_metrics = _scores(y_test, y_pred, y_prob)
            tuned_metrics["threshold"] = thr
            tuning = {
                "model": best_name,
                "best_params": search.best_params_,
                "metrics": tuned_metrics,
                "threshold": thr,
            }
            # si mejora PR-AUC, reemplaza best
            if tuned_metrics["pr_auc"] >= results[best_name]["pr_auc"]:
                best_model = tuned
                best_threshold = thr
                results[best_name].update(tuned_metrics)
                results[best_name]["tuned"] = True
                preds[best_name] = y_pred
                probs[best_name] = y_prob
                thresholds[best_name] = thr
                results[best_name]["figures"]["confusion"] = _save_confusion(
                    y_test, y_pred, best_name
                )

    # McNemar: mejor vs cada otro
    mcnemar = {}
    y_test_arr = y_test.to_numpy() if hasattr(y_test, "to_numpy") else np.asarray(y_test)
    for name in results:
        if name == best_name:
            continue
        try:
            mcnemar[f"{best_name}_vs_{name}"] = mcnemar_test(
                y_test_arr, preds[best_name], preds[name]
            )
        except Exception:
            b = int(np.sum((preds[best_name] == y_test_arr) & (preds[name] != y_test_arr)))
            c = int(np.sum((preds[best_name] != y_test_arr) & (preds[name] == y_test_arr)))
            chi2 = 0.0 if (b + c) == 0 else (abs(b - c) - 1) ** 2 / (b + c)
            mcnemar[f"{best_name}_vs_{name}"] = {
                "chi2": float(chi2),
                "p_value": None,
                "b": b,
                "c": c,
            }

    joblib.dump(
        {
            "model": best_model,
            "features": FEATURE_COLS,
            "model_name": best_name,
            "threshold": best_threshold,
        },
        BEST_MODEL_PATH,
    )

    accuracy_cmp = _save_accuracy_comparison(results)

    payload = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "split": {
            "cut": cut,
            "n_train": int(len(train_df)),
            "n_test": int(len(test_df)),
            "prevalencia_train_pct": round(100 * float(y_train.mean()), 2),
            "prevalencia_test_pct": round(100 * float(y_test.mean()), 2),
            **(
                {k: v for k, v in split_meta.items() if k != "cut"}
                if isinstance(split_meta, dict)
                else {}
            ),
        },
        "models": results,
        "best_model": best_name,
        "best_threshold": best_threshold,
        "cv": {"n_folds": n_folds_eff, "results": cv_table},
        "tuning": tuning,
        "mcnemar": mcnemar,
        "selection_rule": "Mayor PR-AUC; desempate por MCC (umbral óptimo calibrado en train)",
        "figures": {"accuracy_comparison": accuracy_cmp},
        "base_models": [n for n in results if n.startswith("M")],
    }
    METRICS_PATH.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    SPLIT_PATH.write_text(json.dumps(payload["split"], indent=2), encoding="utf-8")
    _write_reports(payload, df)
    return payload


def _write_reports(payload: dict[str, Any], df: pd.DataFrame) -> dict[str, str]:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    # Excel
    xlsx = REPORTS_DIR / "comparativa_modelos.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Comparativa"
    headers = [
        "Modelo",
        "Precision",
        "Recall",
        "F1",
        "Accuracy",
        "ROC-AUC",
        "PR-AUC",
        "MCC",
        "Train_s",
    ]
    ws.append(headers)
    for name, m in payload["models"].items():
        ws.append(
            [
                name,
                m["precision"],
                m["recall"],
                m["f1"],
                m["accuracy"],
                m["roc_auc"],
                m["pr_auc"],
                m["mcc"],
                m["train_seconds"],
            ]
        )
    wb.save(xlsx)

    # Word
    docx_path = REPORTS_DIR / "informe_resultados.docx"
    doc = Document()
    doc.add_heading("Informe demo — Predicción de fallas CMMS", level=1)
    doc.add_paragraph(
        f"Mejor modelo: {payload['best_model']}. "
        f"Regla de selección: {payload['selection_rule']}."
    )
    doc.add_paragraph(
        f"Split: train={payload['split']['n_train']}, test={payload['split']['n_test']}, corte={payload['split']['cut']}."
    )
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Modelo", "PR-AUC", "ROC-AUC", "F1", "MCC"]):
        hdr[i].text = h
    for name, m in payload["models"].items():
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"{m['pr_auc']:.3f}"
        row[2].text = f"{m['roc_auc']:.3f}"
        row[3].text = f"{m['f1']:.3f}"
        row[4].text = f"{m['mcc']:.3f}"
    doc.add_paragraph(
        "Interpretación: se prioriza PR-AUC por el carácter desbalanceado del evento correctivo; "
        "MCC complementa el balance entre clases. Las diferencias pareadas se contrastan con McNemar."
    )
    doc.save(docx_path)

    # PDF
    pdf_path = REPORTS_DIR / "informe_resultados.pdf"
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Informe demo — Predicción de fallas CMMS", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Mejor modelo: <b>{payload['best_model']}</b>", styles["Normal"]),
        Spacer(1, 8),
    ]
    data = [["Modelo", "PR-AUC", "ROC-AUC", "F1", "MCC", "s"]]
    for name, m in payload["models"].items():
        data.append(
            [
                name,
                f"{m['pr_auc']:.3f}",
                f"{m['roc_auc']:.3f}",
                f"{m['f1']:.3f}",
                f"{m['mcc']:.3f}",
                f"{m['train_seconds']:.2f}",
            ]
        )
    t = Table(data, hAlign="LEFT")
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(t)
    story.append(Spacer(1, 12))
    heat = FIGURES_DIR / "heatmap_corr.png"
    if heat.exists():
        story.append(Paragraph("Mapa de calor de correlaciones", styles["Heading2"]))
        story.append(RLImage(str(heat), width=400, height=300))
    best = payload["best_model"]
    roc = FIGURES_DIR / results_fig(payload, best, "roc")
    if roc.exists():
        story.append(Paragraph(f"Curva ROC — {best}", styles["Heading2"]))
        story.append(RLImage(str(roc), width=360, height=300))
    doc_pdf = SimpleDocTemplate(str(pdf_path), pagesize=A4)
    doc_pdf.build(story)

    return {"xlsx": xlsx.name, "docx": docx_path.name, "pdf": pdf_path.name}


def results_fig(payload: dict[str, Any], model: str, kind: str) -> str:
    return payload["models"][model]["figures"][kind]


def load_metrics() -> dict[str, Any] | None:
    if not METRICS_PATH.exists():
        return None
    return json.loads(METRICS_PATH.read_text(encoding="utf-8"))


def predict_fleet(top_n: int = 15) -> list[dict[str, Any]]:
    if not BEST_MODEL_PATH.exists():
        raise FileNotFoundError("No hay modelo entrenado. Ejecuta /api/train primero.")
    bundle = joblib.load(BEST_MODEL_PATH)
    model = bundle["model"]
    features = bundle["features"]
    thr = float(bundle.get("threshold") or 0.5)
    df = _ensure_data()
    latest = df.sort_values("cutoff_date").groupby("vehicle_id", as_index=False).tail(1)
    proba = model.predict_proba(latest[features])[:, 1]
    latest = latest.copy()
    latest["prob_falla_30d"] = proba
    latest["riesgo"] = np.where(
        latest["prob_falla_30d"] >= max(thr, 0.7),
        "alto",
        np.where(latest["prob_falla_30d"] >= thr, "medio", "bajo"),
    )
    out = latest.sort_values("prob_falla_30d", ascending=False).head(top_n)
    cols = [
        "vehicle_id",
        "placa",
        "cutoff_date",
        "prob_falla_30d",
        "riesgo",
        "correctivos_90d",
        "km_30d",
        "dias_desde_ultima_ot",
    ]
    records = out[cols].copy()
    records["cutoff_date"] = records["cutoff_date"].astype(str)
    records["prob_falla_30d"] = records["prob_falla_30d"].round(4)
    return records.to_dict(orient="records")
